from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from script_to_video_production_agent.io_utils import load_bundle_from_source, parse_scenes_from_text, read_text
from script_to_video_production_agent.models import ProjectProfile


class ParsingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.physical_profile = ProjectProfile(
            project_id="fixture-physical",
            title="Replace a Wall-Mounted Air Filter Safely",
            tutorial_type="physical",
            audience="maintenance technicians",
        )

    def test_parse_structured_text(self) -> None:
        text = (
            "Scene 1\n"
            "Narration:\n"
            "\"Lock the unit before service.\"\n"
            "Visual:\n"
            "Technician isolates the power source.\n"
        )
        scenes = parse_scenes_from_text(text, "physical")
        self.assertEqual(len(scenes), 1)
        self.assertEqual(scenes[0].number, 1)
        self.assertEqual(scenes[0].narration[0], "\"Lock the unit before service.\"")
        self.assertEqual(scenes[0].visuals[0], "Technician isolates the power source.")
        self.assertTrue(scenes[0].narration_sha256)

    def test_fallback_paragraph_import(self) -> None:
        text = "First instruction line.\n\nSecond instruction line."
        scenes = parse_scenes_from_text(text, "software")
        self.assertEqual(len(scenes), 2)
        self.assertIn("Show the software interface", scenes[0].visuals[0])

    def test_docx_import(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "script.docx"
            doc = Document()
            doc.add_paragraph("Scene 1")
            doc.add_paragraph("Narration:")
            doc.add_paragraph("\"Confirm the dashboard is empty.\"")
            doc.add_paragraph("Visual:")
            doc.add_paragraph("Operator opens a blank dashboard canvas.")
            doc.save(str(path))
            bundle = load_bundle_from_source(path, self.physical_profile)
            self.assertEqual(bundle.scenes[0].narration[0], "\"Confirm the dashboard is empty.\"")
            self.assertIn("blank dashboard canvas", bundle.scenes[0].visuals[0])
            self.assertIn("Scene 1", read_text(path))


if __name__ == "__main__":
    unittest.main()
